# doc_handler.py
import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

DOC_FILE = "HighlightedNotes.docx"

def init_doc():
    if os.path.exists(DOC_FILE):
        return Document(DOC_FILE)
    else:
        doc = Document()
        doc.add_heading("Smart Vocabulary Notes", 0)
        doc.save(DOC_FILE)
        return doc

def add_sentence_to_doc(sentence, new_words):
    doc = init_doc()
    para = doc.add_paragraph()
    words = sentence.split()

    for w in words:
        run = para.add_run(w + " ")
        if w.lower().strip(".,?!") in [nw.lower() for nw in new_words]:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.bold = True

    doc.save(DOC_FILE)
