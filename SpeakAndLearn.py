import os
import re
import datetime
import subprocess
import tkinter as tk
from tkinter import messagebox
from openpyxl import Workbook, load_workbook
from docx import Document
from docx.shared import RGBColor

# --- File paths ---
EXCEL_FILE = "SmartVocabularyNotes.xlsx"
DOC_FILE = "HighlightedNotes.docx"
SHEET_NAME = "New Words"
ICON_FILE = "icon.ico"  # Must be in the same folder

# ✅ Create or load Excel workbook
def get_workbook():
    if os.path.exists(EXCEL_FILE):
        wb = load_workbook(EXCEL_FILE)
        if SHEET_NAME in wb.sheetnames:
            ws = wb[SHEET_NAME]
        else:
            ws = wb.create_sheet(SHEET_NAME)
            ws.append(["No.", "New Words", "Sentence", "Date/Time"])
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = SHEET_NAME
        ws.append(["No.", "New Words", "Sentence", "Date/Time"])
    return wb, ws

# ✅ Extract new words ignoring [brackets]
def extract_new_words(sentence, existing_words):
    cleaned = re.sub(r"\[.*?\]", "", sentence)
    words = re.findall(r"\b[a-zA-Z']+\b", cleaned)
    new_words = []
    for w in words:
        lw = w.lower()
        if lw not in existing_words:
            existing_words.add(lw)
            new_words.append(w)
    return new_words

# ✅ Load existing words from Excel (Column 2)
def load_existing_words(ws):
    words = set()
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[1]:
            words.add(str(row[1]).lower())
    return words

# ✅ Update DOC file with highlighted new words
def update_doc_file(sentence_no, sentence, new_words):
    if os.path.exists(DOC_FILE):
        doc = Document(DOC_FILE)
    else:
        doc = Document()
        doc.add_heading("Highlighted Notes", level=1)

    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{sentence_no}. ")

    for w in sentence.split():
        clean_w = re.sub(r"[^\w\s]", "", w)
        if clean_w in new_words:
            run = paragraph.add_run(w + " ")
            run.font.bold = True
            run.font.highlight_color = 7  # Yellow
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            paragraph.add_run(w + " ")

    doc.save(DOC_FILE)

# ✅ Add sentence
def add_sentence():
    sentence = entry.get().strip()
    if not sentence:
        return

    wb, ws = get_workbook()
    existing_words = load_existing_words(ws)
    new_words = extract_new_words(sentence, existing_words)
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sentence_no = ws.max_row

    if new_words:
        for w in new_words:
            ws.append([sentence_no, w, sentence, timestamp])
            sentence_no += 1
    else:
        ws.append([sentence_no, "", sentence, timestamp])
        sentence_no += 1

    wb.save(EXCEL_FILE)
    update_doc_file(sentence_no - 1, sentence, [w for w in new_words])

    entry.delete(0, tk.END)
    messagebox.showinfo("Saved", f"Sentence saved. {len(new_words)} new words highlighted in DOC.")

# ✅ Open Excel file
def open_excel():
    if os.path.exists(EXCEL_FILE):
        subprocess.Popen(["start", EXCEL_FILE], shell=True)
    else:
        messagebox.showwarning("Not Found", "Excel file not found.")

# ✅ Open DOC file
def open_doc():
    if os.path.exists(DOC_FILE):
        subprocess.Popen(["start", DOC_FILE], shell=True)
    else:
        messagebox.showwarning("Not Found", "DOC file not found.")

# ✅ GUI
def start_gui():
    root = tk.Tk()
    root.title("Smart Vocabulary Notes")
    root.geometry("520x230")
    tk.Label(root, text="Type your sentence:").pack(pady=5)
    global entry
    entry = tk.Entry(root, width=60)
    entry.pack(pady=5)

    tk.Button(root, text="Add Sentence", command=add_sentence).pack(pady=8)

    # ✅ Open file buttons
    btn_frame = tk.Frame(root)
    btn_frame.pack(pady=5)
    tk.Button(btn_frame, text="📊 View Excel", command=open_excel, width=15).grid(row=0, column=0, padx=5)
    tk.Button(btn_frame, text="📝 View Doc", command=open_doc, width=15).grid(row=0, column=1, padx=5)

    root.mainloop()

if __name__ == "__main__":
    start_gui()
