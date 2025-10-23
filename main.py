"""
LingoBaby – Smart Vocabulary (v3 Optimized)
-------------------------------------------
✅ Uses JSON-based irregular verbs
✅ Auto-learns new irregulars (no external libs)
✅ Saves to Excel + Word
✅ Works offline and with PyInstaller
"""

import tkinter as tk
from tkinter import messagebox
import re, os, subprocess, json
from datetime import datetime
import inflect
import nltk
from nltk.corpus import wordnet
from openpyxl import Workbook, load_workbook
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from update_if_irregular_v2 import add_if_irregular

# -----------------------------
# ⚙️ Initialization
# -----------------------------
os.environ["TYPEGUARD_DISABLE"] = "1"
nltk.download("wordnet", quiet=True)

p = inflect.engine()
EXCEL_FILE = "SmartVocabularyNotes.xlsx"
DOC_FILE = "HighlightedNotes.docx"
SHEET_NAME = "New Words"
ICON_FILE = "LingoBaby.ico"
IRREGULAR_JSON = "irregular_verbs_extended.json"


# -----------------------------
# 🧩 Irregular Verbs Loader
# -----------------------------
def load_irregulars():
    """Load irregular verbs from JSON file once."""
    if not os.path.exists(IRREGULAR_JSON):
        print(f"⚠️ JSON not found: {IRREGULAR_JSON}")
        return {}, {}
    try:
        with open(IRREGULAR_JSON, "r", encoding="utf-8") as f:
            data = json.load(f)
        verbs = {v["base"].lower(): [v["past"], v["past_participle"], v["ing"], v["s"]] for v in data}
        mapping = {form.lower(): base.lower() for base, forms in verbs.items() for form in forms}
        return verbs, mapping
    except Exception as e:
        print(f"⚠️ Failed to read irregular verbs JSON: {e}")
        return {}, {}

IRREGULAR_VERBS, IRREGULAR_MAP = load_irregulars()


# -----------------------------
# 📘 Excel Helpers
# -----------------------------
def init_workbook():
    """Initialize Excel file with header if missing."""
    if not os.path.exists(EXCEL_FILE):
        wb = Workbook()
        ws = wb.active
        ws.title = SHEET_NAME
        ws.append(["No.", "New Word", "Sentence", "Explanation", "Date/Time"])
        wb.save(EXCEL_FILE)
        return wb, ws
    wb = load_workbook(EXCEL_FILE)
    return wb, wb[SHEET_NAME]


def get_existing_words():
    """Return set of all saved words."""
    try:
        wb, ws = init_workbook()
        return {str(row[1]).lower() for row in ws.iter_rows(min_row=2, values_only=True) if row[1]}
    except Exception:
        return set()


def add_new_sentence(words, sentence):
    """Append new words to Excel."""
    wb, ws = init_workbook()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    next_row = ws.max_row + 1
    for w in words:
        ws.append([next_row, w, sentence, get_explanation(w), now])
        next_row += 1
    wb.save(EXCEL_FILE)


# -----------------------------
# 📝 Word Document Helpers
# -----------------------------
def init_doc():
    if not os.path.exists(DOC_FILE):
        doc = Document()
        doc.add_heading("Highlighted Vocabulary Notes", level=1)
        doc.save(DOC_FILE)
    return Document(DOC_FILE)


def get_next_num(doc):
    count = sum(1 for p in doc.paragraphs if re.match(r"^\d+\.", p.text.strip()))
    return count + 1


def add_sentence_to_doc(sentence, words):
    doc = init_doc()
    num = get_next_num(doc)
    p = doc.add_paragraph()
    p.add_run(f"{num}. ").bold = True

    for word in sentence.split():
        clean = re.sub(r"[^\w']", "", word)
        base = get_base_form(clean)
        run = p.add_run(word + " ")
        if base.lower() in [w.lower() for w in words]:
            run.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(DOC_FILE)


# -----------------------------
# 📚 Dictionary / Explanation
# -----------------------------
def get_base_form(word: str) -> str:
    w = word.lower().strip()
    if w in IRREGULAR_VERBS:
        return w
    if w in IRREGULAR_MAP:
        return IRREGULAR_MAP[w]
    for suf in ["ing", "ed", "es", "s"]:
        if w.endswith(suf) and len(w) > len(suf) + 1:
            base = w[:-len(suf)]
            if wordnet.synsets(base):
                return base
    return w


def get_explanation(word: str) -> str:
    """Return explanation with correct grammatical logic (noun/verb detection)."""
    word = word.lower().strip()
    parts = []

    # Get synsets and their parts of speech
    synsets = wordnet.synsets(word)
    pos_tags = {s.pos() for s in synsets}  # e.g., {'n', 'v', 'a'}

    # 1️⃣ If it's a noun → show plural
    if "n" in pos_tags:
        plural = p.plural(word)
        if plural and plural != word:
            parts.append(f"Plural: {plural}")

    # 2️⃣ If it's a verb → show irregular or regular forms
    if "v" in pos_tags or word in IRREGULAR_VERBS or word in IRREGULAR_MAP:
        base = get_base_form(word)
        if base in IRREGULAR_VERBS:
            forms = ", ".join(IRREGULAR_VERBS[base])
            parts.append(f"Verb forms: {base}, {forms}")
        else:
            ing = word + "ing" if not word.endswith("ing") else word
            past = word + "ed" if not word.endswith("ed") else word
            third = word + "s" if not word.endswith("s") else word
            parts.append(f"Verb forms: {word}, {ing}, {past}, {third}")

    # 3️⃣ Meaning (always shown if available)
    if synsets:
        meaning = synsets[0].definition()
        parts.append(f"Meaning: {meaning}")

    # If nothing matched, fallback meaning only
    if not parts:
        parts.append("Meaning: (No definition found)")

    return " | ".join(parts)



# -----------------------------
# 🧠 Word Extraction
# -----------------------------
def extract_new_words(sentence, existing):
    tokens = re.findall(r"\b[a-zA-Z']+\b", sentence)
    return [get_base_form(t) for t in tokens if get_base_form(t).lower() not in existing]


# -----------------------------
# 🔍 Open / Search / Add
# -----------------------------
def open_excel():
    if os.path.exists(EXCEL_FILE):
        subprocess.Popen(["start", EXCEL_FILE], shell=True)
    else:
        messagebox.showwarning("Not found", "Excel file not found.")


def open_doc():
    if os.path.exists(DOC_FILE):
        subprocess.Popen(["start", DOC_FILE], shell=True)
    else:
        messagebox.showwarning("Not found", "Doc file not found.")


def search_word():
    word = entry.get().strip().lower()
    if not word:
        messagebox.showwarning("Empty", "Enter a word to search.")
        return

    base = get_base_form(word)
    if not os.path.exists(EXCEL_FILE):
        messagebox.showwarning("Missing", "Excel not found.")
        return

    wb = load_workbook(EXCEL_FILE)
    ws = wb[SHEET_NAME]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[1] and str(row[1]).lower() == base:
            messagebox.showinfo(
                "Found",
                f"✅ '{base}' found!\n\n📖 Sentence: {row[2]}\n💬 Explanation: {row[3]}\n🕓 Added: {row[4]}",
            )
            return
    messagebox.showinfo("Not Found", f"❌ '{word}' not in your list.")


def on_submit():
    sentence = entry.get().strip()
    if not sentence:
        messagebox.showwarning("Empty", "Please enter a sentence.")
        return

    existing = get_existing_words()
    new_words = extract_new_words(sentence, existing)

    if not new_words:
        messagebox.showinfo("Info", "No new words found.")
        return

    for w in new_words:
        try:
            add_if_irregular(w)
        except Exception as e:
            print(f"(⚠️ Skipped irregular check for '{w}': {e})")

    add_new_sentence(new_words, sentence)
    add_sentence_to_doc(sentence, new_words)
    messagebox.showinfo("Success", f"Added: {', '.join(new_words)}")
    entry.delete(0, tk.END)


# -----------------------------
# 🪟 GUI
# -----------------------------
def start_gui():
    root = tk.Tk()
    root.title("LingoBaby – Smart Vocabulary")
    root.geometry("660x360")

    if os.path.exists(ICON_FILE):
        try:
            root.iconbitmap(ICON_FILE)
        except Exception:
            pass

    tk.Label(root, text="Enter a sentence or word:", font=("Segoe UI", 11, "bold")).pack(pady=8)
    global entry
    entry = tk.Entry(root, width=75)
    entry.pack(pady=5)

    frame = tk.Frame(root)
    frame.pack(pady=12)
    tk.Button(frame, text="Add Sentence", width=18, command=on_submit).grid(row=0, column=0, padx=6)
    tk.Button(frame, text="Search Word", width=18, command=search_word).grid(row=0, column=1, padx=6)
    tk.Button(frame, text="📊 View Excel", width=18, command=open_excel).grid(row=1, column=0, padx=6)
    tk.Button(frame, text="📝 View Doc", width=18, command=open_doc).grid(row=1, column=1, padx=6)

    tk.Label(root, text="Tip: Use base words (e.g. 'go', not 'goes').", fg="gray").pack(pady=8)
    root.mainloop()


# -----------------------------
# 🚀 Run
# -----------------------------
if __name__ == "__main__":
    start_gui()
