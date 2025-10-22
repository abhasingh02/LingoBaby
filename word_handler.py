# word_handler.py
import os
import re
from docx import Document
from docx.shared import RGBColor, Pt
from config import DOC_FILE
from helpers import clean_word_for_compare, close_word_if_open

def init_document():
    """Close Word if open, then load or create DOCX."""
    close_word_if_open()
    if os.path.exists(DOC_FILE):
        doc = Document(DOC_FILE)
    else:
        doc = Document()
        doc.add_heading("Highlighted Notes", level=1)
    return doc

def get_next_sentence_number(doc):
    """
    Find the highest leading number in existing paragraphs like 'N. ' and return next number.
    If none found, return 1.
    """
    max_n = 0
    import re
    for p in doc.paragraphs:
        text = p.text.strip()
        m = re.match(r"^\s*(\d+)\.", text)
        if m:
            try:
                n = int(m.group(1))
                if n > max_n:
                    max_n = n
            except Exception:
                pass
    return max_n + 1

def append_sentence(doc, sentence_no, sentence, new_words_set):
    """
    Append one numbered paragraph to the document. Bold + highlight the new words.
    new_words_set should contain cleaned (lowercase) words to highlight.
    """
    p = doc.add_paragraph()
    p.add_run(f"{sentence_no}. ")
    for token in sentence.split():
        # Determine word cleaned (ignore bracketed tokens for comparison)
        cleaned = re.sub(r"[^\w\s]", "", token)
        cleaned_key = clean_word_for_compare(cleaned)
        if cleaned_key and cleaned_key in new_words_set:
            r = p.add_run(token + " ")
            r.bold = True
            # Word highlight via python-docx uses WD_COLOR_INDEX - but here we use run.font.highlight_color attribute not available
            # Using a combination: bold + set font color to black (default) and rely on bold to emphasize
            # However python-docx supports highlight by run.font.highlight_color = WD_COLOR_INDEX.YELLOW if available
            try:
                from docx.enum.text import WD_COLOR_INDEX
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            except Exception:
                pass
            r.font.color.rgb = RGBColor(0, 0, 0)
            r.font.size = Pt(12)
        else:
            r = p.add_run(token + " ")
            r.font.size = Pt(12)

def save_document(doc):
    doc.save(DOC_FILE)
